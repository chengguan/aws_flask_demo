from docx import Document
import sys
import getopt

def print_usage():
    print('Usage:')
    print(f'    python {__file__} -i <inputfile>')

def load_docx(file_name):
    with open(file_name, 'rb') as f:
        document = Document(f)
    return document

def make_wiki_table(table):
    cols = len(table.columns)
    rows = len(table.rows)

    wikitable = ''

    wikitable += "{| class='wikitable' style='text-align: center;'\n"
    for r in range(0, rows):
        wikitable += '|-\n'
        for c in range(0, cols):
            try:
                write_cell = (c == table.cell(r, c)._tc.left) & (r == table.cell(r, c)._tc.top)
                col_span = table.cell(r, c)._tc.right - c
                row_span = table.cell(r, c)._tc.bottom - r
            except:
                write_cell = True
                col_span = 1
                row_span = 1
            if write_cell:
                wikitable +=  f'| rowspan={row_span}, colspan={col_span} | {table.cell(r, c).text}\n'
    wikitable +=  "|}\n\n"

    return wikitable

def print_tc(cell):
    print(f'Text: {cell.text} TOP:{cell._tc.top} BOT:{cell._tc.bottom} L:{cell._tc.left} R:{cell._tc.right}')

def make_wikitable_from_file(filename):
    print(f"Loading {filename}")
    doc = load_docx(filename)

    # Display the able in the wikitable form.
    print(f"Number of tables found: {len(doc.tables)}")
    i = 1
    wiki_tables = []
    for each_table in doc.tables:
        print(f'Table {i} is {len(each_table.columns)} x {len(each_table.rows)}')
        i += 1
        wiki_tables.append(make_wiki_table(each_table))

    return wiki_tables

def main(argv):
    inputfile = ''
    
    try:
        # setup the option flags
        opts, args = getopt.getopt(argv, "hi:", ["ifile="])
    except getopt.GetoptError:
        print_usage()
        sys.exit(2)
    for opt, arg in opts:
        if opt == '-h':
            print_usage()
            sys.exit()
        elif opt in ("-i", "--ifile"):
            inputfile = arg

    print(f'Input file is {inputfile}')

    # Load the document
    doc = load_docx(inputfile)

    # Display the able in the wikitable form.
    print(f"Number of tables found: {len(doc.tables)}")
    i = 1
    for each_table in doc.tables:
        print(f'Table {i} is {len(each_table.columns)} x {len(each_table.rows)}')
        i += 1
        make_wiki_table(each_table)

if __name__ == "__main__":
    main(sys.argv[1:])